from docx import Document

def write_to_docx(text, file_path):
    try:
        document = Document(file_path)
    except:
        document = Document()
    
    document.add_paragraph(text)
    document.save(file_path)   


def modifyArr(arr):
    arr[1][0] *= arr[0][1]
    arr[1][2] *= arr[0][1]
    arr[1][0] -= arr[0][0] * arr[1][1]
    arr[1][2] -= arr[0][2] * arr[1][1]
    arr[1][1] = 0
    return arr
import numpy as np

def Gauss(A, F):
    n = len(A)
    a = np.concatenate((A, F), axis = 1)  #делаем расширенную матрицу 
    for i in range (n):             #прямой ход
        a[i] /= a[i][i]
        for j in range (i + 1, n):
            a[j] -= a[i] * a[j][i]

    for i in range (n - 1, -1, -1):  #обратный ход
        for j in range (i - 1, -1, -1):
            a[j] -= a[i] * a[j][i]

    ans = []
    for i in range (n):
        ans.append(a[i][n])
    return ans

def minusOneMod(exp, fi, filename):
    write_to_docx('---промежуточный расчет {} ^-1 mod {} ---'.format(exp, fi), filename)
    x = pow(int(exp), -1, fi)
    write_to_docx(f'ПОЛУЧИЛИ: {x}---промежуточный расчет закончен---', filename)
    return x

def Gauss_finite(A, F, p, filename):
    n = len(A)
    write_to_docx("Применяем алгоритм Гаусса для решения, прямой ход:", filename)
    a = np.hstack((A,F)) #делаем расширенную матрицу 
    write_to_docx(str(a), filename)
    for i in range (n):          #прямой ход
        write_to_docx(f"a[{i}][{i}]^-1 = {minusOneMod(a[i][i], p, filename)}", filename)
        a[i] *= minusOneMod(a[i][i], p, filename)
        a[i] = a[i] % p
        write_to_docx("После умножения коэффециенты такие:" + str(a[i]), filename)
        for j in range (i + 1, n):
            write_to_docx(f"Вычитаем из {j}-й строчки {i}-ю", filename)
            a[j] -= a[i] * a[j][i]
            a[j] = a[j] % p
            write_to_docx(f"После вычитания из {j}-й строчки {i}-й: коэффициенты следующие ", filename)
            write_to_docx(str(a[j]), filename)


    write_to_docx("Обратный ход: ", filename)
    for i in range (n - 1, -1, -1):  #обратный ход
        for j in range (i - 1, -1, -1):
            # print("i,j", i,j)
            write_to_docx(f"После подставления в {j}-ю строчку {i}-й: коэффициенты следующие ", filename)
            a[j] -= a[i] * a[j][i]
            a[j] = a[j] % p
            write_to_docx(str(a[j]), filename)
    # print(a)

    ans = []
    for i in range (n):
        ans.append(-a[i][n] % p)
    return ans


def blackly(p, numbers_in_sled, two_dimension, filename):
    a = numbers_in_sled
    arr = two_dimension #np.asarray
    fs = np.zeros(a-1)
    fs = fs.reshape(a-1,1)
    write_to_docx('Берем {} любые суперплоскости из следов (все вычисления в F{}):'.format(a - 1, p), filename)
    for i in range(0, a - 1):
        st = ''
        for j in range (0, a):
            if j != a - 1:
                st += ('{} * x{} + '.format(arr[i][j], j + 1))
            if j == a - 1:
                st += (str(arr[i][j]))
        write_to_docx(st + ' = 0', filename)

    write_to_docx('Выразив коэффициенты получаем:', filename)
    ans = Gauss_finite(arr, fs, p, filename)
    for i in range(a-1):
        write_to_docx(f'x{i+1} = {ans[i]}', filename)
    return f'{ans}\nПервое число - ответ'

# blackly(11, 4, [[2, 4, 10, 8], [8, 2, 2, 3], [3,4,2,8]], "blackly_test.docx")