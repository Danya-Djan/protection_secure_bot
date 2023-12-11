from docx import Document

def write_to_docx(text, file_path):
    try:
        document = Document(file_path)
    except:
        document = Document()
    
    document.add_paragraph(text)
    document.save(file_path)  

def minusOneMod(exp, fi, filename):
    write_to_docx('---промежуточный расчет {} ^-1 mod {} ---'.format(exp, fi), filename)
    x = pow(exp, -1, fi)
    write_to_docx(f'ПОЛУЧИЛИ: {x}---промежуточный расчет закончен---', filename)
    return x

def Blom(p, Ax, Ay, AxClosed, AyClosed, Bx, By, BxClosed, ByClosed, filename):
    write_to_docx('Матрица доверенного центра имеет вид \n(a  b)\n(b  c)\n', filename)
    write_to_docx('YA = ({}) = (a  b)({})\n     ({}) = (b  c)({})'.format(AxClosed, Ax, AyClosed, Ay), filename)
    write_to_docx('YB = ({}) = (a  b)({})\n     ({}) = (b  c)({})'.format(BxClosed, Bx, ByClosed, By), filename)
    write_to_docx('Откуда получаем систему уравнений', filename)
    write_to_docx('{}a + {}b = {}'.format(Ax, Ay, AxClosed), filename)
    write_to_docx('{}b + {}c = {}'.format(Ax, Ay, AyClosed), filename)
    write_to_docx('{}a + {}b = {}'.format(Bx, By, BxClosed), filename)
    write_to_docx('Решив систему уравнений получаем', filename)
    arr = [[Ax, Ay, -AxClosed],[Ax, Ay, -AyClosed],[Bx, By, -BxClosed]]
    arr[2][1] *=  arr[0][0]
    arr[2][2] *=  arr[0][0]
    arr[2][1] += arr[2][0]*(-arr[0][1])
    arr[2][2] += arr[2][0]*(-arr[0][2])
    tChisl = pow(-arr[2][2], 1, p)
    tZnam = minusOneMod(pow(arr[2][1], 1, p), p, filename)
    b = pow(tChisl*tZnam, 1, p)

    arr = [[Ax, Ay, -AxClosed],[Ax, Ay, -AyClosed],[Bx, By, -BxClosed]]
    tChisl = pow(-arr[0][2] - arr[0][1]*b, 1, p)
    tZnam = minusOneMod(pow(arr[0][0], 1, p), p, filename)
    a = pow(tChisl*tZnam, 1, p)

    arr = [[Ax, Ay, -AxClosed],[Ax, Ay, -AyClosed],[Bx, By, -BxClosed]]
    tChisl = pow(-arr[1][2] - arr[1][0]*b, 1, p)
    tZnam = minusOneMod(pow(arr[1][1], 1, p), p, filename)
    c = pow(tChisl*tZnam, 1, p)

    write_to_docx('a = {} / {} mod {} = {}*{} mod {} = {}'.format(-arr[0][2] - arr[0][1]*b, arr[0][0], p, tChisl, tZnam, p, a), filename)
    write_to_docx('b = {} / {} mod {} = {}*{} mod {} = {}'.format(-arr[2][2], arr[2][1], p, tChisl,tZnam, p, b), filename)
    write_to_docx('c = {} / {} mod {} = {}*{} mod {} = {}'.format(-arr[1][2] - arr[1][0]*b, arr[1][1], p, tChisl, tZnam, p, c), filename)

    write_to_docx('Откуда получаем матрицу доверительного центра', filename)
    write_to_docx('D = ({}  {})\n    ({}  {})'.format(a, b, b, c), filename)

    s = pow(AxClosed*Bx + AyClosed*By, 1, p)
    write_to_docx('Общий сеансовый ключ S = Sa = Sb = (закр А)(откр B) = ({},{})({},{})(<- 2ю вертикально) = {} mod {} = {}'
      .format(AxClosed, AyClosed, Bx, By, AxClosed*Bx + AyClosed*By, p, s), filename)
    
# Blom(11, 8, 4, 10, 5 , 2, 5, 2, 10, "blom_test.docx")