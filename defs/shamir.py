from docx import Document

def write_to_docx(text, file_path):
    try:
        document = Document(file_path)
    except:
        document = Document()
    
    document.add_paragraph(text)
    document.save(file_path)   

def minusOneMod(exp, fi, filename):
    write_to_docx('---промежуточный расчет {} ^-1 mod {} ---'.format(exp, fi), filename)
    x = pow(exp, -1, fi)
    write_to_docx(f'ПОЛУЧИЛИ: {x}---промежуточный расчет закончен---', filename)
    return x

def Shamir(p, x1, y1, x2, y2, x3, y3, filename):
    arr = []
    write_to_docx('l_1(x) = (x - x2)/(x1 - x2) * (x - x3)/(x1 - x3) = (x - {})/({} - {}) * (x - {})/({} - {}) ='
      .format(x2, x1, x2, x3, x1, x3), filename)
    znam = pow((x1- x2)*(x1-x3), 1, p)
    a = 1
    b = pow(-x2 -x3, 1, p)
    c = pow((-x2)*(-x3),1 , p)
    write_to_docx('=1/{} * (({})x^2 + ({})x + {}) ='.format(znam,a, b, c), filename)
    znam = minusOneMod(znam, p, filename)
    a = pow(a*znam, 1, p)
    b = pow(b*znam, 1, p)
    c = pow(c*znam, 1, p)
    write_to_docx('=(({})x^2 + ({})x + {}'.format(a, b, c), filename)
    arr.append([a, b, c])

    write_to_docx('l_2(x) = (x - x1)/(x2 - x1) * (x - x3)/(x2 - x3) = (x - {})/({} - {}) * (x - {})/({} - {}) ='
      .format(x1, x2, x1, x3, x2, x3), filename)
    znam = pow((x2 - x1)*(x2 - x3), 1, p)
    a = 1
    b = pow(-x1 -x3, 1, p)
    c = pow((-x1)*(-x3),1 , p)
    write_to_docx('=1/{} * (({})x^2 + ({})x + {}) ='.format(znam,a, b, c), filename)
    znam = minusOneMod(znam, p, filename)
    a = pow(a*znam, 1, p)
    b = pow(b*znam, 1, p)
    c = pow(c*znam, 1, p)
    write_to_docx('=(({})x^2 + ({})x + {}'.format(a, b, c), filename)
    arr.append([a, b, c])

    write_to_docx('l_3(x) = (x - x1)/(x3 - x1) * (x - x2)/(x3 - x2) = (x - {})/({} - {}) * (x - {})/({} - {}) ='
      .format(x1, x3, x1, x2, x3, x2), filename)
    znam = pow((x3 - x1)*(x3 - x2), 1, p)
    a = 1
    b = pow(-x1 -x2, 1, p)
    c = pow((-x1)*(-x2),1 , p)
    write_to_docx('=1/{} * (({})x^2 + ({})x + {}) ='.format(znam,a, b, c), filename)
    znam = minusOneMod(znam, p, filename)
    a = pow(a*znam, 1, p)
    b = pow(b*znam, 1, p)
    c = pow(c*znam, 1, p)
    write_to_docx('=(({})x^2 + ({})x + {}'.format(a, b, c), filename)
    arr.append([a, b, c])

    a = pow(arr[0][0]*y1 + arr[1][0]*y2 + arr[2][0]*y3, 1, p)
    write_to_docx('a = {}y1 + {}y2 + {}y3 = {}*{} + {}*{} + {}*{} = {} mod {} = {}'
      .format(arr[0][0], arr[1][0], arr[2][0], arr[0][0], y1, arr[1][0], y2, arr[2][0], y3, arr[0][0]*y1 + arr[1][0]*y2 + arr[2][0]*y3, p , a), filename)

    b = pow(arr[0][1]*y1 + arr[1][1]*y2 + arr[2][1]*y3, 1, p)
    write_to_docx('b = {}y1 + {}y2 + {}y3 = {}*{} + {}*{} + {}*{} = {} mod {} = {}'
      .format(arr[0][1], arr[1][1], arr[2][1], arr[0][1], y1, arr[1][1], y2, arr[2][1], y3, arr[0][1]*y1 + arr[1][1]*y2 + arr[2][1]*y3, p , b), filename)

    M = pow(arr[0][2]*y1 + arr[1][2]*y2 + arr[2][2]*y3, 1, p)
    write_to_docx('M = {}y1 + {}y2 + {}y3 = {}*{} + {}*{} + {}*{} = {} mod {} = {}'
      .format(arr[0][2], arr[1][2], arr[2][2], arr[0][2], y1, arr[1][2], y2, arr[2][2], y3, arr[0][2]*y1 + arr[1][2]*y2 + arr[2][2]*y3, p , M), filename)

    write_to_docx('Ответ будет вида ax^2 + bx + m = {}x^2 + {}x + {} <= последний член - секрет'.format(a, b, M), filename)
    return f'{a}, {b}, {M} <- секрет'


# Shamir(11, 4, 3, 5, 8, 6, 1, "shamir_test.docx")
