from docx import Document

def write_to_docx(text, file_path):
    try:
        document = Document(file_path)
    except:
        document = Document()
    
    document.add_paragraph(text)
    document.save(file_path)  
    
def group_points(a, b, n, file_name):

    l = [[0 for j in range(8)] for i in range(n)]

    Result = []
    write_to_docx(f"Задача: уравнение эллиптической кривой y^2 = x^3 + ({a})*x + ({b}); F_{n}", file_name)
    write_to_docx("Решение: получение группы точек с помощью таблицы:", file_name)

    for i, item in enumerate(l):
        l[i][0] = i
        l[i][1] = pow(i**2, 1, n)
        l[i][2] = pow(i**3, 1, n)
        l[i][3] = pow(l[i][0] * a, 1, n)
        l[i][4] = b
        l[i][5] = pow((l[i][2] + l[i][3] + l[i][4]), 1, n)

    #находим y1 y2
    for i, item in enumerate(l):
        for j in range (0, n):
            if l[i][5] == l[j][1]:
                if l[i][6] == 0:
                    l[i][6] = l[j][0]
                else:
                    l[i][7] = l[j][0]

    #находим точки
    for i, item in enumerate(l):
        if l[i][6] != 0 and l[i][7] != 0:
            Result.append([l[i][0], l[i][6]])
            Result.append([l[i][0], l[i][7]])

    Result.append(0)
    formatted_text = ""
    write_to_docx('x    x^2  x^3  {}x{}{}{}y^2 y1   y2'.format(a, ' '*(4 - len(str(a))), b, ' '*(5 - len(str(b)))), file_name)
    for row in l:
        row_text = ""
        for elem in row:
            row_text += f"{elem}{' ' * (5 - len(str(elem)))}"
        formatted_text += f"{row_text}\n"
        
    write_to_docx(formatted_text, file_name)
        
    write_to_docx('Точки: ' + str(Result), file_name)
    write_to_docx('Размер группы точек:' + str(len(Result)), file_name)
    
group_points(-1, -14, 17, "group_points_test.docx")