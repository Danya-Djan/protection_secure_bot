from docx import Document

def write_to_docx(text, file_path):
    try:
        document = Document(file_path)
    except:
        document = Document()
    
    document.add_paragraph(text)
    document.save(file_path)   

def module(a):
    if a >= 0:
        return a
    else:
        return a - 2 * a

def findDotMainFunc(x, y, a, f, c, filename):
    resX = x
    resY = y
    l = []
    l.append([x, y])
    for i in range(0, c):
        write_to_docx('Рассчет {}А:'.format(i + 1), filename)
        write_to_docx('{}A + A = ({},{}) + ({},{})'.format(i + 1, resX, resY, x, y), filename)
        if i == 0:
            p = module(a)
            ch = pow(3 * resX ** 2 - p, 1, f)
            write_to_docx('(3x^2 - p) mod F = (3*{}^2 + ({}) )mod {} = {}'.format(x, a, f, ch), filename)
            zn = minusOneMod(pow(2 * resY, 1, f), f, filename)
            res = pow(ch * zn, 1, f)
            write_to_docx('lambda = ((3x^2 + а )/ 2y) mod F = ({} * {}) mod {} = {}'.format(ch, zn, f, res), filename)
            resXPrev = resX
            resYPrev = resY
            resX = pow(res ** 2 - 2 * x, 1, f)
            resY = pow(-y + res * (x - resX), 1, f)
            write_to_docx('X{} = lambda^2 -2x mod F = {} mod {} = {}'.format(i + 1, res ** 2 - 2 * x, f, resX), filename)
            write_to_docx('Y{} = -y + lambda*(X - X{}) mod F = {} mod {} = {}'.format(i + 1, i + 1,-y + res * (x - resX), f, resY), filename)
            l.append([resX, resY])
        elif i >= 1:
            if resX == x:
                write_to_docx('x2 = x1 => деление на 0', filename)
                write_to_docx('Значит, -Y0 == Y{} mod {}'.format(i, f), filename)
                l.append(0)
                break
            ch = pow(resY - y, 1, f)
            write_to_docx('y{} - y{} mod F = ({} - {}) mod {} = {}'.format(i + 1, i, resY, y, f, ch), filename)
            zn = minusOneMod(pow(resX - x, 1, f), f, filename)
            write_to_docx('(x{} - x{})^-1 mod F = ({} - {})^-1 mod {} = {}'.format(i + 1, i, resX, x, f, zn), filename)
            res = pow(ch * zn, 1, f)
            write_to_docx('lambda = ({} * {}) mod {} = {}'.format(ch, zn, f, res), filename)
            resXPrev = resX
            resYPrev = resY
            resX = pow(res ** 2 - resX - x, 1, f)
            resY = pow(-resY + res * (resXPrev - resX), 1, f)
            write_to_docx('X{} = (lambda^2 - x{} - x{}) mod F = ({} - ({}) - ({}))mod {} = {}'
                  .format(i + 1, i, x, res ** 2, resXPrev, x, f, resX), filename)
            write_to_docx('Y{} = (-y{} + lambda*(X{} - X{})) mod F = ({} + {}) mod {} = {}'
                  .format(i + 1, i, i , i+1, -resYPrev, res * (resXPrev - resX), f, resY), filename)
            l.append([resX, resY])
        write_to_docx('{}A + A = ({},{}) + ({},{}) = ({},{})'.format(i + 1,resXPrev, resYPrev, x, y,  resX, resY), filename)
    return l

def minusOneMod(exp, fi, filename):
    write_to_docx('---промежуточный расчет {} ^-1 mod {} ---'.format(exp, fi), filename)
    x = pow(exp, -1, fi)
    write_to_docx(f'ПОЛУЧИЛИ: {x}---промежуточный расчет закончен---', filename)
    return x



def diffi2(generator1, generator2, a, b, F, alisa, bob, filename):
    write_to_docx('A = a*G = {}*({} {})\n'.format(alisa, generator1, generator2), filename)
    vec = findDotMainFunc(generator1, generator2, a, F, alisa - 1, filename)
    vec = vec[len(vec) - 1]
    write_to_docx('\nГенерация общего сеансового ключа s', filename)
    write_to_docx('S = b*A = {}*({} {})\n'.format(bob, vec[0], vec[1]), filename)
    vec = findDotMainFunc(vec[0],vec[1], a, F, bob - 1, filename)
    vec = vec[len(vec) - 1]
    write_to_docx('\nОбщий сеансовый ключ s =' + str(vec), filename)
    return vec

# diffi2(6, 16, -3, -10, 17, 2, 3, 'diffi2_test.docx')