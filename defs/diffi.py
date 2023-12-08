from docx import Document

def DecToBin(a):
    t = a
    bin = []
    print('Перведем число',a,'в двоичный вид:')
    while a != 0:
        if a % 2 == 0:
            bin.append(0)
            print(a, '0')
            a /= 2
        if a % 2 == 1:
            print(a, '1')
            bin.append(1)
            a //= 2
    bin = bin[::-1]
    print(t,'(x10) = (x2) ',bin)
    bin = bin[::-1]
    return bin

def write_to_docx(text, file_path):
    try:
        document = Document(file_path)
    except:
        document = Document()
    
    document.add_paragraph(text)
    document.save(file_path)   
    
def FastExponentiation(num1, num2, num3):
    power = bin(num2)
    print(power)
    power = power[2:]
    print(num2, 'в двоичной:', power)
    r1 = 1
    r2 = num1
    r3 = 1
    for i in range(len(power ) - 1):
        r3 = r1 * (r2 ** int(power[i])) % num3
        print(r1, '*',(r2 ** int(power[i])), '=', r3)
        r1 = r3 ** 2 % num3
        print(r3, '^2', '=', r1)
    answer = r1 * (r2 ** int(power[len(power ) - 1])) % num3
    print(r1, '*', r2 ** int(power[len(power ) - 1]), '= ', answer)
    print('Ответ:', answer)
    return answer
def diffi_public(g, p, a, b, filename):

  aCounter = 1
  bCounter = 1
  
  while pow(g, aCounter, p) != a:
    write_to_docx(f"g^{aCounter} mod {p} == {pow(g, aCounter, p)}", filename)
    aCounter += 1
  write_to_docx('A = {} = g^a mod p = {}^a mod {}'.format(a, g, p), filename)
  write_to_docx(f'Перебором получаем, что a = {aCounter}', filename)
  while pow(g, bCounter, p) != b:
    write_to_docx(f"g^{bCounter} mod {p} == {pow(g, bCounter, p)}", filename)
    bCounter += 1
  write_to_docx('B = {} = g^b mod p = {}^b mod {}'.format(b, g, p), filename)
  write_to_docx(f'Перебором получаем, что b = {bCounter}', filename)

  s1 = pow(b, aCounter, p)
  s2 = pow(a, bCounter, p)
  write_to_docx('Сеансовый ключ s = B^a mod p = {}^{} mod {} = {}'.format(b, aCounter, p, s1), filename)
  write_to_docx('Сеансовый ключ s = A^b mod p = {}^{} mod {} = {}'.format(a, bCounter, p, s2), filename)

def diffi_private(g, p, a, b, filename):
  A = pow(g, a, p)
  B = pow(g, b, p)

  write_to_docx('A = {} = g^a mod p = {}^a mod {}'.format(pow(g, a, p), g, p), filename)

  write_to_docx('B = {} = g^b mod p = {}^b mod {}'.format(pow(g, b, p), g, p), filename)

  s1 = pow(B, a, p)
  s2 = pow(A, b, p)
  write_to_docx('Сеансовый ключ s = B^a mod p = {}^{} mod {} = {}'.format(B, a, p, s1), filename)
  write_to_docx('Сеансовый ключ s = A^b mod p = {}^{} mod {} = {}'.format(A, b, p, s2), filename)
  
diffi_public(14, 17, 13, 10, "diffi_test.docx")