from docx import Document

def write_to_docx(text, file_path):
    try:
        document = Document(file_path)
    except:
        document = Document()
    
    document.add_paragraph(text)
    document.save(file_path)  
    

def module(a):
    if a >= 0:
        return a
    else:
        return a - 2 * a

def MinusOneMod(exp, fi, filename):
    write_to_docx('---промежуточный расчет {} ^-1 mod {} ---'.format(exp, fi), filename)
    x = pow(exp, -1, fi)
    write_to_docx(f'ПОЛУЧИЛИ: {x}---промежуточный расчет закончен---', filename)
    return x


def gostfindDotMainFunc(whileX, whileY, resX, resY, x, y, f, filename):
    i = 2
    while ((whileX != resX) and (whileY != resY)):
        if resX == x:
            write_to_docx('x2 = x1 => деление на 0')
            break
        ch = pow(resY - y, 1, f)
        write_to_docx('y2 - y1 mod F = ({} - {}) mod {} = {}'.format(resY, y, f, ch), filename)
        zn = MinusOneMod(pow(resX - x, 1, f), f, filename)
        write_to_docx('(x2 - x1)^-1 mod F = ({} - {})^-1 mod {} = {}'.format(resX, x, f, zn), filename)
        res = pow(ch * zn, 1, f)
        write_to_docx('lambda = ({} * {}) mod {} = {}'.format(ch, zn, f, res), filename)
        resXPrev = resX
        resYPrev = resY
        resX = pow(res ** 2 - resX - x, 1, f)
        resY = pow(-resY + res * (resXPrev - resX), 1, f)
        write_to_docx('Xз = (lamda^2 - x2 - x1) mod F = ({} + ({}) + ({}))mod {} = {}'
            .format(res ** 2, resXPrev, x, f, resX), filename)
        write_to_docx('Yз = (-y + landa*(X - Xз)) mod F = ({} + {}) mod {} = {}'
            .format(-resY, res * (resXPrev - resX), f, resY), filename)
        write_to_docx('Промежуточный результат: ({},{})'.format(resX, resY), filename)
        i += 1
        if i >= 25:
            write_to_docx('Что-то i какое-то большое, скорее всего не является генератором...', filename)
            break
    write_to_docx(f'Откуда получаем d ={i}', filename)
    return i
def findDotMainFunc(x, y, a, f, c, filename):
    resX = x
    resY = y
    l = []
    l.append([x, y])
    for i in range(0, c):
        write_to_docx('Рассчет {}А:'.format(i + 1), filename)
        write_to_docx('{}A + A = ({},{}) + ({},{})'.format(i + 1, resX, resY, x, y), filename)
        if i == 0:
            p = module(a)
            ch = pow(3 * resX ** 2 - p, 1, f)
            write_to_docx('(3x^2 - p) mod F = (3*{}^2 + ({}) )mod {} = {}'.format(x, a, f, ch), filename)
            zn = MinusOneMod(pow(2 * resY, 1, f), f, filename)
            res = pow(ch * zn, 1, f)
            write_to_docx('lambda = ((3x^2 + а )/ 2y) mod F = ({} * {}) mod {} = {}'.format(ch, zn, f, res), filename)
            resXPrev = resX
            resYPrev = resY
            resX = pow(res ** 2 - 2 * x, 1, f)
            resY = pow(-y + res * (x - resX), 1, f)
            write_to_docx('X{} = lambda^2 -2x mod F = {} mod {} = {}'.format(i + 1, res ** 2 - 2 * x, f, resX), filename)
            write_to_docx('Y{} = -y + lambda*(X - X{}) mod F = {} mod {} = {}'.format(i + 1, i + 1,-y + res * (x - resX), f, resY), filename)
            l.append([resX, resY])
        elif i >= 1:
            if resX == x:
                write_to_docx('x2 = x1 => деление на 0', filename)
                write_to_docx('Значит, -Y0 == Y{} mod {}'.format(i, f), filename)
                l.append(0)
                break
            ch = pow(resY - y, 1, f)
            write_to_docx('y{} - y{} mod F = ({} - {}) mod {} = {}'.format(i + 1, i, resY, y, f, ch), filename)
            zn = MinusOneMod(pow(resX - x, 1, f), f, filename)
            write_to_docx('(x{} - x{})^-1 mod F = ({} - {})^-1 mod {} = {}'.format(i + 1, i, resX, x, f, zn), filename)
            res = pow(ch * zn, 1, f)
            write_to_docx('lambda = ({} * {}) mod {} = {}'.format(ch, zn, f, res), filename)
            resXPrev = resX
            resYPrev = resY
            resX = pow(res ** 2 - resX - x, 1, f)
            resY = pow(-resY + res * (resXPrev - resX), 1, f)
            write_to_docx('X{} = (lambda^2 - x{} - x{}) mod F = ({} - ({}) - ({}))mod {} = {}'
                  .format(i + 1, i, x, res ** 2, resXPrev, x, f, resX), filename)
            write_to_docx('Y{} = (-y{} + lambda*(X{} - X{})) mod F = ({} + {}) mod {} = {}'
                  .format(i + 1, i, i , i+1, -resYPrev, res * (resXPrev - resX), f, resY), filename)
            l.append([resX, resY])
        write_to_docx('{}A + A = ({},{}) + ({},{}) = ({},{})'.format(i + 1,resXPrev, resYPrev, x, y,  resX, resY), filename)
    return l


def gost(m, a, b, f, g_1, g_2, n, q_1, q_2, k, filename):
    
    x = g_1
    y = g_2
    e = q_1
    d = q_2

    p = module(a)
    write_to_docx(f'Из у-я ЭП: p = {p}', filename)
    write_to_docx('С = k*G = {}*G ='.format(k), filename)

    c = findDotMainFunc(x, y, a, f, k - 1, filename)
    c = c[len(c) - 1]
    write_to_docx(f'Откуда с ={c}', filename)

    write_to_docx('\nНайдем закрытый ключ:\nq = d*G = ({},{}) = d*({},{})'.format(e, d, x, y), filename)
    i = 1
    temp1 = x
    temp2 = y
    write_to_docx(f'{d},{e}', filename)
    while(e != temp1 or d != temp2):
        ans = findDotMainFunc(x,y , a, f, i, filename)
        temp1 = ans[len(ans) - 1][0]
        temp2 = ans[len(ans) - 1][1]
        i += 1
    d = i
    write_to_docx(f"d = {d}", filename)
    Sy = pow(c[0]*d + k * m ,1, n)
    write_to_docx('\nНайдем эл.подпись s: (Sy = ) (Xc*d + k*m) mod n = ({}*{} + {}*{}) mod {} = {}'
        .format(c[0], d, k, m, n, Sy), filename)
    write_to_docx('\nОтвет: ({},{})'.format(c[0], Sy), filename)
    
    final = f'({c[0]}, {Sy})'
    return final
    
# gost(5, -8, -5, 11, 8, 5, 15, 2, 3, 2, "gost_test.docx")